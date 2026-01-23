import uuid
import zipfile
from pathlib import Path
from typing import Optional, Any, Dict, List
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from reportlab.lib.colors import HexColor
from app.config import settings
from app.schemas.profile import ProfileSchema
from app.schemas.tailor import TailorResponse, SuggestedBullet
from app.utils.encryption import encryption_service
from app.utils.app_logger import app_logger
import json
import os


class ResumeTemplate:
    """Resume template configuration"""
    def __init__(self, template_id: str, name: str, description: str, style: str = "modern"):
        self.template_id = template_id
        self.name = name
        self.description = description
        self.style = style


class DocumentGenerator:
    """Generate tailored resume documents (DOCX and PDF)"""

    def __init__(self):
        # Convert to absolute path (handles relative paths like ./data/uploads)
        upload_path = Path(settings.storage_local_dir)
        if not upload_path.is_absolute():
            # Make relative to the API directory
            api_dir = Path(__file__).parent.parent.parent  # Go up from services -> app -> api
            upload_path = api_dir / upload_path
        self.upload_dir = upload_path.resolve()
        self.upload_dir.mkdir(parents=True, exist_ok=True)

        # Create templates directory
        self.templates_dir = upload_path.parent / "templates"
        self.templates_dir.mkdir(parents=True, exist_ok=True)

        # Available resume templates with enhanced descriptions
        self.templates = {
            "modern-professional": ResumeTemplate(
                "modern-professional",
                "Modern Professional",
                "Clean, contemporary design with professional typography, perfect for tech, marketing, and business roles",
                "modern"
            ),
            "executive-classic": ResumeTemplate(
                "executive-classic",
                "Executive Classic",
                "Traditional layout with serif fonts and formal structure, ideal for senior management and executive positions",
                "classic"
            ),
            "creative-minimal": ResumeTemplate(
                "creative-minimal",
                "Creative Minimal",
                "Minimalist design with creative accents, perfect for designers, creatives, and startup environments",
                "minimal"
            ),
            "tech-focused": ResumeTemplate(
                "tech-focused",
                "Tech Focused",
                "Optimized for technical roles with code-friendly formatting, ideal for developers and engineers",
                "tech"
            ),
            "academic-formal": ResumeTemplate(
                "academic-formal",
                "Academic Formal",
                "Formal layout suitable for academic and research positions, with emphasis on publications and credentials",
                "academic"
            )
        }

    def _parse_description_into_bullets(self, description: str) -> list[str]:
        """Parse description text into multiple bullet points using intelligent splitting strategies"""
        if not description:
            return []

        # First, split by existing line breaks
        lines = [line.strip() for line in description.split('\n') if line.strip()]

        # If there are multiple lines, treat each as a bullet
        if len(lines) > 1:
            return lines

        # If it's one long paragraph, split by sentences and logical breaks
        text = description.strip()

        # Strategy 1: Split by periods followed by capital letters (most reliable)
        import re
        sentences = re.split(r'\.\s*(?=[A-Z])', text)
        sentences = [s.strip() for s in sentences if s.strip()]
        if len(sentences) > 1:
            return [s + ('.' if not s.endswith('.') else '') for s in sentences if len(s) > 1]

        # Strategy 2: Split by any periods
        sentences = [s.strip() for s in text.split('.') if s.strip()]
        if len(sentences) > 1:
            return [s + ('.' if not s.endswith('.') else '') for s in sentences if len(s) > 1]

        # Strategy 3: Split by semicolons
        sentences = [s.strip() for s in text.split(';') if s.strip()]
        if len(sentences) > 1:
            return sentences

        # Strategy 4: Split by commas (if more than 2)
        comma_parts = [s.strip() for s in text.split(',') if s.strip()]
        if len(comma_parts) > 2:
            return comma_parts

        # Strategy 5: Force split long text into chunks of ~100 characters
        if len(text) > 200:
            chunks = []
            remaining = text
            while len(remaining) > 100:
                chunk_end = 100
                # Try to find a good break point near the end
                for i in range(min(100, len(remaining) - 1), 50, -1):
                    if remaining[i] in [' ', ',']:
                        chunk_end = i
                        break
                chunks.append(remaining[:chunk_end].strip())
                remaining = remaining[chunk_end:].strip()
            if remaining:
                chunks.append(remaining)
            return [chunk for chunk in chunks if len(chunk) > 10]

        # If no good splitting points found, return the whole text as one bullet
        return [text]

    def get_available_templates(self) -> Dict[str, Dict[str, str]]:
        """Get list of available templates"""
        return {
            template_id: {
                "name": template.name,
                "description": template.description,
                "style": template.style
            }
            for template_id, template in self.templates.items()
        }

    async def generate_tailored_resume(
        self,
        profile: ProfileSchema,
        tailoring: TailorResponse,
        job_posting: Any,
        template_id: str = "modern-professional",
        original_profile_id: Optional[str] = None,
        original_file_id: Optional[str] = None,
        original_file_ext: Optional[str] = None
    ) -> tuple[str, Optional[str], Optional[str]]:
        """Generate tailored application documents, return file IDs

        Returns: (docx_resume_id, cover_letter_docx_id, application_package_id)

        Uses pure Python libraries:
        - python-docx for DOCX generation
        """

        # Get selected template
        template = self.templates.get(template_id, self.templates["modern-professional"])

        # Generate documents using selected template
        docx_file_id = await self._generate_docx_with_template(profile, tailoring, template, original_profile_id)

        # Generate cover letter DOCX (always from template)
        cover_letter_file_id = await self._generate_cover_letter_docx(profile, tailoring, original_profile_id)

        # Generate application package ZIP (containing resume and cover letter)
        # Create empty ZIP for now - will be populated when downloaded
        package_file_id = str(uuid.uuid4())
        zip_path = self.upload_dir / f"{package_file_id}_application_package.zip"
        app_logger.log_info(f"Creating ZIP at: {zip_path}, upload_dir: {self.upload_dir}")

        # Create an empty ZIP file as placeholder
        with zipfile.ZipFile(str(zip_path), 'w', zipfile.ZIP_DEFLATED) as zip_file:
            # Add a placeholder file
            zip_file.writestr("README.txt", "Application package will be generated on download.")
        app_logger.log_info(f"Created placeholder ZIP: {package_file_id}")

        return docx_file_id, cover_letter_file_id, package_file_id

    async def _generate_from_original(
        self,
        profile: ProfileSchema,
        tailoring: TailorResponse,
        original_file_id: str,
        original_file_ext: str,
        original_profile_id: Optional[str]
    ) -> tuple[str, Optional[str]]:
        """Generate tailored resume by modifying the original uploaded file"""

        # Find the original file
        original_file_path = None
        if original_file_ext == ".docx":
            original_file_path = self.upload_dir / f"{original_file_id}.docx"
        elif original_file_ext == ".pdf":
            original_file_path = self.upload_dir / f"{original_file_id}.pdf"
        else:
            # Unsupported format, fall back to template
            return await self._generate_docx_template(profile, tailoring, original_profile_id), None

        if not original_file_path or not original_file_path.exists():
            # Original file not found, fall back to template
            return await self._generate_docx_template(profile, tailoring, original_profile_id), None

        # Generate new file IDs for the tailored versions
        tailored_docx_id = str(uuid.uuid4())
        tailored_pdf_id = str(uuid.uuid4())

        try:
            if original_file_ext == ".docx":
                # Modify existing DOCX file
                tailored_docx_path = self.upload_dir / f"{tailored_docx_id}.docx"
                tailored_pdf_path = self.upload_dir / f"{tailored_pdf_id}.pdf"

                # Copy and modify the original DOCX
                await self._modify_docx_file(str(original_file_path), str(tailored_docx_path), profile, tailoring)

                # Generate PDF from the modified DOCX
                pdf_id = await self._generate_pdf_from_docx(str(tailored_docx_path), tailored_pdf_id)
                if pdf_id is None:
                    return tailored_docx_id, None

                return tailored_docx_id, tailored_pdf_id

            elif original_file_ext == ".pdf":
                # For PDFs, we need to convert to editable format, modify, then convert back
                # This is more complex, so for now fall back to template
                return await self._generate_docx_template(profile, tailoring, original_profile_id), None

        except Exception as e:
            app_logger.log_error(f"Failed to modify original file {original_file_id}: {e}")
            # Fall back to template generation
            return await self._generate_docx_template(profile, tailoring, original_profile_id), None

    async def _modify_docx_file(
        self,
        original_path: str,
        output_path: str,
        profile: ProfileSchema,
        tailoring: TailorResponse
    ) -> None:
        """Modify an existing DOCX file with tailored content"""
        try:
            # Copy the original file
            import shutil
            shutil.copy2(original_path, output_path)

            # Open and modify the document
            doc = Document(output_path)

            # For now, we'll add tailored content at the end
            # In a more sophisticated implementation, we could try to find and replace
            # specific sections, but that requires advanced text analysis

            # Add a section break and tailored content
            doc.add_page_break()

            # Add tailored work experience
            if tailoring.suggested_bullets and len(tailoring.suggested_bullets) > 0:
                doc.add_heading('Tailored Experience Highlights', level=1)

                for bullet in tailoring.suggested_bullets:
                    if bullet.tailored:
                        para = doc.add_paragraph()
                        para.add_run(f"• {bullet.tailored}")

            # Add tailored summary if available
            if tailoring.jd_summary:
                doc.add_heading('Position Summary', level=1)
                summary_para = doc.add_paragraph()
                summary_para.add_run(tailoring.jd_summary)

            # Save the modified document
            doc.save(output_path)
            app_logger.log_info(f"Modified DOCX file saved: {output_path}")

        except Exception as e:
            app_logger.log_error(f"Failed to modify DOCX file: {e}")
            # If modification fails, just copy the original
            import shutil
            shutil.copy2(original_path, output_path)

    async def _generate_docx_template(
        self,
        profile: ProfileSchema,
        tailoring: TailorResponse,
        original_profile_id: Optional[str]
    ) -> str:
        """Fallback: Generate DOCX using template approach"""
        return await self._generate_docx(profile, tailoring, original_profile_id)

    async def _generate_pdf_from_docx(
        self,
        docx_path: str,
        pdf_id: str
    ) -> Optional[str]:
        """Generate PDF from DOCX file"""
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.units import inch

            pdf_path = self.upload_dir / f"{pdf_id}.pdf"

            # For now, create a simple PDF with tailored content
            # In a production system, you'd use a proper DOCX to PDF converter
            doc = SimpleDocTemplate(str(pdf_path), pagesize=letter)
            styles = getSampleStyleSheet()
            story = []

            # Add title
            title_style = styles['Heading1']
            story.append(Paragraph("Tailored Resume", title_style))
            story.append(Spacer(1, 0.25*inch))

            # Add content placeholder
            content_style = styles['Normal']
            story.append(Paragraph("Your tailored resume content has been prepared.", content_style))

            doc.build(story)
            app_logger.log_info(f"PDF generated from DOCX: {pdf_id}")
            return pdf_id

        except Exception as e:
            app_logger.log_error(f"PDF generation from DOCX failed: {e}")
            return None

    async def _generate_docx_with_template(
        self,
        profile: ProfileSchema,
        tailoring: TailorResponse,
        template: ResumeTemplate,
        original_profile_id: Optional[str]
    ) -> str:
        """Generate tailored resume DOCX using selected template"""
        doc = Document()

        # Apply template-specific styling
        if template.style == "modern":
            await self._apply_modern_template(doc, profile, tailoring)
        elif template.style == "classic":
            await self._apply_classic_template(doc, profile, tailoring)
        elif template.style == "minimal":
            await self._apply_minimal_template(doc, profile, tailoring)
        elif template.style == "tech":
            await self._apply_tech_template(doc, profile, tailoring)
        elif template.style == "academic":
            await self._apply_academic_template(doc, profile, tailoring)
        else:
            # Default to modern
            await self._apply_modern_template(doc, profile, tailoring)

        # Save document
        file_id = str(uuid.uuid4())
        file_path = self.upload_dir / f"{file_id}.docx"
        doc.save(str(file_path))
        app_logger.log_info(f"DOCX generated with {template.name} template: {file_id}")
        return file_id

    async def _apply_modern_template(self, doc: Document, profile: ProfileSchema, tailoring: TailorResponse):
        """Apply modern professional template - clean, contemporary design"""
        # Set professional margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        # ===== HEADER SECTION =====
        # Name - Large, bold, centered
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(f"{profile.basics.firstName} {profile.basics.lastName}")
        name_run.font.size = Pt(24)
        name_run.bold = True
        name_run.font.name = 'Calibri'
        name_run.font.color.rgb = RGBColor(33, 37, 41)
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Professional title
        if profile.basics.headline:
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(profile.basics.headline)
            title_run.font.size = Pt(14)
            title_run.font.name = 'Calibri'
            title_run.font.color.rgb = RGBColor(0, 123, 255)  # Professional blue
            title_run.bold = True
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Contact information bar
        contact_parts = []
        if profile.basics.phone:
            contact_parts.append(f"📞 {profile.basics.phone}")
        if profile.basics.email:
            contact_parts.append(f"✉ {profile.basics.email}")
        if profile.basics.location:
            contact_parts.append(f"📍 {profile.basics.location}")

        # Add professional links
        if profile.basics.links and len(profile.basics.links) > 0:
            for link in profile.basics.links:
                if link.url:
                    if link.type == "LinkedIn":
                        contact_parts.append(f"💼 {link.label or 'LinkedIn'}")
                    elif link.type == "GitHub":
                        contact_parts.append(f"💻 {link.label or 'GitHub'}")
                    elif link.type == "Portfolio":
                        contact_parts.append(f"🎨 {link.label or 'Portfolio'}")
                    else:
                        contact_parts.append(f"🔗 {link.label or link.type}")

        if contact_parts:
            contact_para = doc.add_paragraph(" | ".join(contact_parts))
            contact_run = contact_para.runs[0]
            contact_run.font.size = Pt(10)
            contact_run.font.name = 'Calibri'
            contact_run.font.color.rgb = RGBColor(108, 117, 125)
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add professional separator line
        separator = doc.add_paragraph()
        separator_run = separator.add_run("─" * 60)
        separator_run.font.color.rgb = RGBColor(0, 123, 255)
        separator.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # ===== PROFESSIONAL SUMMARY =====
        if profile.basics.summary:
            doc.add_paragraph()  # Spacing

            # Section header with accent
            summary_header = doc.add_paragraph()
            summary_header_run = summary_header.add_run("PROFESSIONAL SUMMARY")
            summary_header_run.font.size = Pt(12)
            summary_header_run.bold = True
            summary_header_run.font.name = 'Calibri'
            summary_header_run.font.color.rgb = RGBColor(0, 123, 255)
            summary_header.paragraph_format.space_after = Pt(6)

            # Summary content
            summary_para = doc.add_paragraph()
            summary_run = summary_para.add_run(profile.basics.summary)
            summary_run.font.size = Pt(10)
            summary_run.font.name = 'Calibri'
            summary_run.font.color.rgb = RGBColor(73, 80, 87)

        # ===== WORK EXPERIENCE =====
        if profile.work_experience:
            doc.add_paragraph()  # Section spacing

            # Section header
            exp_header = doc.add_paragraph()
            exp_header_run = exp_header.add_run("WORK EXPERIENCE")
            exp_header_run.font.size = Pt(14)
            exp_header_run.bold = True
            exp_header_run.font.name = 'Calibri'
            exp_header_run.font.color.rgb = RGBColor(0, 123, 255)
            exp_header.paragraph_format.space_after = Pt(8)

            for exp in profile.work_experience:
                # Company name - prominent
                company_para = doc.add_paragraph()
                company_run = company_para.add_run(exp.company)
                company_run.font.size = Pt(13)
                company_run.bold = True
                company_run.font.name = 'Calibri'
                company_run.font.color.rgb = RGBColor(33, 37, 41)
                company_para.paragraph_format.space_after = Pt(2)

                # Job title and dates
                title_date_para = doc.add_paragraph()
                title_run = title_date_para.add_run(exp.title)
                title_run.font.size = Pt(11)
                title_run.font.name = 'Calibri'
                title_run.font.color.rgb = RGBColor(52, 58, 64)

                # Add dates and location
                date_location = ""
                if exp.startDate:
                    date_location = exp.startDate
                    if exp.endDate:
                        date_location += f" - {exp.endDate}"
                    else:
                        date_location += " - Present"
                if exp.location:
                    date_location += f" | {exp.location}"

                if date_location:
                    title_date_para.add_run(f"  |  {date_location}")
                    date_run = title_date_para.runs[-1]
                    date_run.font.size = Pt(9)
                    date_run.font.color.rgb = RGBColor(108, 117, 125)

                # Job description with professional bullet points
                if exp.description:
                    description_bullets = self._parse_description_into_bullets(exp.description)
                    for bullet in description_bullets:
                        bullet_para = doc.add_paragraph()
                        bullet_run = bullet_para.add_run(f"• {bullet}")
                        bullet_run.font.size = Pt(10)
                        bullet_run.font.name = 'Calibri'
                        bullet_run.font.color.rgb = RGBColor(73, 80, 87)
                        bullet_para.paragraph_format.left_indent = Inches(0.25)
                        bullet_para.paragraph_format.space_after = Pt(3)

                # Technologies - subtle formatting
                if exp.technologies and len(exp.technologies) > 0:
                    tech_para = doc.add_paragraph()
                    tech_label = tech_para.add_run("Technologies: ")
                    tech_label.font.size = Pt(9)
                    tech_label.font.name = 'Calibri'
                    tech_label.bold = True
                    tech_label.font.color.rgb = RGBColor(0, 123, 255)

                    tech_content = tech_para.add_run(", ".join(exp.technologies))
                    tech_content.font.size = Pt(9)
                    tech_content.font.name = 'Calibri'
                    tech_content.font.color.rgb = RGBColor(108, 117, 125)

                doc.add_paragraph()  # Spacing between jobs

        # ===== EDUCATION =====
        if profile.education:
            doc.add_paragraph()

            edu_header = doc.add_paragraph()
            edu_header_run = edu_header.add_run("EDUCATION")
            edu_header_run.font.size = Pt(14)
            edu_header_run.bold = True
            edu_header_run.font.name = 'Calibri'
            edu_header_run.font.color.rgb = RGBColor(0, 123, 255)

            for edu in profile.education:
                school_para = doc.add_paragraph()
                school_run = school_para.add_run(edu.school)
                school_run.font.size = Pt(12)
                school_run.bold = True
                school_run.font.name = 'Calibri'
                school_run.font.color.rgb = RGBColor(33, 37, 41)

                degree_para = doc.add_paragraph()
                degree_text = f"{edu.degree}"
                if edu.field:
                    degree_text += f" in {edu.field}"
                degree_run = degree_para.add_run(degree_text)
                degree_run.font.size = Pt(10)
                degree_run.font.name = 'Calibri'
                degree_run.font.color.rgb = RGBColor(52, 58, 64)

                # Dates
                if edu.startDate or edu.endDate:
                    date_para = doc.add_paragraph()
                    date_text = f"{edu.startDate or ''} - {edu.endDate or ''}".strip(" -")
                    date_run = date_para.add_run(date_text)
                    date_run.font.size = Pt(9)
                    date_run.font.name = 'Calibri'
                    date_run.font.color.rgb = RGBColor(108, 117, 125)

                doc.add_paragraph()

        # ===== SKILLS =====
        if profile.skills:
            doc.add_paragraph()

            skills_header = doc.add_paragraph()
            skills_header_run = skills_header.add_run("SKILLS")
            skills_header_run.font.size = Pt(14)
            skills_header_run.bold = True
            skills_header_run.font.name = 'Calibri'
            skills_header_run.font.color.rgb = RGBColor(0, 123, 255)

            # Group skills by category
            skills_by_category = {}
            for skill in profile.skills:
                category = skill.category or 'Technical Skills'
                if category not in skills_by_category:
                    skills_by_category[category] = []
                skill_text = skill.name
                if skill.level:
                    skill_text += f" ({skill.level})"
                skills_by_category[category].append(skill_text)

            for category, skills in skills_by_category.items():
                cat_para = doc.add_paragraph()
                cat_run = cat_para.add_run(f"{category}: ")
                cat_run.font.size = Pt(10)
                cat_run.bold = True
                cat_run.font.name = 'Calibri'
                cat_run.font.color.rgb = RGBColor(52, 58, 64)

                skills_run = cat_para.add_run(", ".join(skills))
                skills_run.font.size = Pt(10)
                skills_run.font.name = 'Calibri'
                skills_run.font.color.rgb = RGBColor(73, 80, 87)

    async def _apply_classic_template(self, doc: Document, profile: ProfileSchema, tailoring: TailorResponse):
        """Apply executive classic template - traditional, formal design with serif fonts"""
        # Set formal margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)

        # ===== FORMAL HEADER =====
        # Name in serif font, centered
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(f"{profile.basics.firstName} {profile.basics.lastName}")
        name_run.font.size = Pt(22)
        name_run.bold = True
        name_run.font.name = 'Times New Roman'
        name_run.font.color.rgb = RGBColor(0, 0, 0)
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_para.paragraph_format.space_after = Pt(8)

        # Professional title
        if profile.basics.headline:
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(profile.basics.headline)
            title_run.font.size = Pt(12)
            title_run.font.name = 'Times New Roman'
            title_run.font.color.rgb = RGBColor(89, 89, 89)
            title_run.italic = True
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Formal contact information
        contact_parts = []
        if profile.basics.phone:
            contact_parts.append(profile.basics.phone)
        if profile.basics.email:
            contact_parts.append(profile.basics.email)
        if profile.basics.location:
            contact_parts.append(profile.basics.location)

        if contact_parts:
            contact_para = doc.add_paragraph(", ".join(contact_parts))
            contact_run = contact_para.runs[0]
            contact_run.font.size = Pt(10)
            contact_run.font.name = 'Times New Roman'
            contact_run.font.color.rgb = RGBColor(89, 89, 89)
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Elegant separator
        separator = doc.add_paragraph()
        separator_run = separator.add_run("• • •")
        separator_run.font.size = Pt(12)
        separator_run.font.name = 'Times New Roman'
        separator_run.font.color.rgb = RGBColor(128, 128, 128)
        separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
        separator.paragraph_format.space_before = Pt(12)
        separator.paragraph_format.space_after = Pt(12)

        # ===== PROFESSIONAL EXPERIENCE =====
        if profile.work_experience:
            exp_header = doc.add_paragraph()
            exp_header_run = exp_header.add_run("PROFESSIONAL EXPERIENCE")
            exp_header_run.font.size = Pt(14)
            exp_header_run.bold = True
            exp_header_run.font.name = 'Times New Roman'
            exp_header_run.font.color.rgb = RGBColor(0, 0, 0)
            exp_header.paragraph_format.space_after = Pt(12)

            for exp in profile.work_experience:
                # Company name - formal styling
                company_para = doc.add_paragraph()
                company_run = company_para.add_run(exp.company)
                company_run.font.size = Pt(12)
                company_run.bold = True
                company_run.font.name = 'Times New Roman'
                company_run.font.color.rgb = RGBColor(0, 0, 0)

                # Job title
                title_para = doc.add_paragraph()
                title_run = title_para.add_run(exp.title)
                title_run.font.size = Pt(11)
                title_run.font.name = 'Times New Roman'
                title_run.font.color.rgb = RGBColor(64, 64, 64)
                title_run.italic = True

                # Dates and location - formal format
                if exp.startDate or exp.location:
                    date_para = doc.add_paragraph()
                    date_parts = []
                    if exp.startDate:
                        date_str = exp.startDate
                        if exp.endDate:
                            date_str += f" - {exp.endDate}"
                        else:
                            date_str += " - Present"
                        date_parts.append(date_str)
                    if exp.location:
                        date_parts.append(exp.location)

                    date_run = date_para.add_run(", ".join(date_parts))
                    date_run.font.size = Pt(10)
                    date_run.font.name = 'Times New Roman'
                    date_run.font.color.rgb = RGBColor(89, 89, 89)

                # Professional bullet points
                if exp.description:
                    description_bullets = self._parse_description_into_bullets(exp.description)
                    for bullet in description_bullets:
                        bullet_para = doc.add_paragraph()
                        bullet_run = bullet_para.add_run(f"• {bullet}")
                        bullet_run.font.size = Pt(10)
                        bullet_run.font.name = 'Times New Roman'
                        bullet_run.font.color.rgb = RGBColor(64, 64, 64)
                        bullet_para.paragraph_format.left_indent = Inches(0.3)

                doc.add_paragraph()

        # ===== EDUCATION =====
        if profile.education:
            edu_header = doc.add_paragraph()
            edu_header_run = edu_header.add_run("EDUCATION")
            edu_header_run.font.size = Pt(14)
            edu_header_run.bold = True
            edu_header_run.font.name = 'Times New Roman'
            edu_header_run.font.color.rgb = RGBColor(0, 0, 0)
            edu_header.paragraph_format.space_before = Pt(18)

            for edu in profile.education:
                school_para = doc.add_paragraph()
                school_run = school_para.add_run(edu.school)
                school_run.font.size = Pt(12)
                school_run.bold = True
                school_run.font.name = 'Times New Roman'
                school_run.font.color.rgb = RGBColor(0, 0, 0)

                degree_para = doc.add_paragraph()
                degree_text = edu.degree
                if edu.field:
                    degree_text += f", {edu.field}"
                degree_run = degree_para.add_run(degree_text)
                degree_run.font.size = Pt(11)
                degree_run.font.name = 'Times New Roman'
                degree_run.font.color.rgb = RGBColor(64, 64, 64)

                # Graduation date
                if edu.endDate:
                    date_para = doc.add_paragraph()
                    date_run = date_para.add_run(f"Graduated: {edu.endDate}")
                    date_run.font.size = Pt(10)
                    date_run.font.name = 'Times New Roman'
                    date_run.font.color.rgb = RGBColor(89, 89, 89)

                doc.add_paragraph()

    async def _apply_minimal_template(self, doc: Document, profile: ProfileSchema, tailoring: TailorResponse):
        """Apply creative minimal template - clean, spacious, creative touches"""
        # Set generous margins for breathing room
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.6)
            section.right_margin = Inches(0.6)

        # ===== MINIMALIST HEADER =====
        # Clean name, left-aligned
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(f"{profile.basics.firstName} {profile.basics.lastName}")
        name_run.font.size = Pt(20)
        name_run.bold = True
        name_run.font.name = 'Helvetica'
        name_run.font.color.rgb = RGBColor(23, 23, 23)

        # Minimal contact bar
        contact_parts = []
        if profile.basics.email:
            contact_parts.append(profile.basics.email)
        if profile.basics.phone:
            contact_parts.append(profile.basics.phone)
        if profile.basics.location:
            contact_parts.append(profile.basics.location)

        if contact_parts:
            contact_para = doc.add_paragraph(" | ".join(contact_parts))
            contact_run = contact_para.runs[0]
            contact_run.font.size = Pt(9)
            contact_run.font.name = 'Helvetica'
            contact_run.font.color.rgb = RGBColor(128, 128, 128)

        # Subtle accent line
        accent = doc.add_paragraph()
        accent_run = accent.add_run("▬")
        accent_run.font.size = Pt(14)
        accent_run.font.color.rgb = RGBColor(0, 184, 148)  # Teal accent
        accent.paragraph_format.space_before = Pt(12)
        accent.paragraph_format.space_after = Pt(12)

        # ===== EXPERIENCE SECTION =====
        if profile.work_experience:
            exp_header = doc.add_paragraph()
            exp_header_run = exp_header.add_run("EXPERIENCE")
            exp_header_run.font.size = Pt(16)
            exp_header_run.bold = True
            exp_header_run.font.name = 'Helvetica'
            exp_header_run.font.color.rgb = RGBColor(0, 184, 148)  # Teal accent

            for exp in profile.work_experience:
                doc.add_paragraph()  # Spacing

                # Company - clean and prominent
                company_para = doc.add_paragraph()
                company_run = company_para.add_run(exp.company)
                company_run.font.size = Pt(14)
                company_run.bold = True
                company_run.font.name = 'Helvetica'
                company_run.font.color.rgb = RGBColor(23, 23, 23)

                # Role and dates on same line
                role_date_para = doc.add_paragraph()
                role_run = role_date_para.add_run(exp.title)
                role_run.font.size = Pt(11)
                role_run.font.name = 'Helvetica'
                role_run.font.color.rgb = RGBColor(64, 64, 64)

                # Add dates and location
                if exp.startDate or exp.location:
                    meta_parts = []
                    if exp.startDate:
                        date_str = exp.startDate
                        if exp.endDate:
                            date_str += f"–{exp.endDate}"
                        else:
                            date_str += "–Present"
                        meta_parts.append(date_str)
                    if exp.location:
                        meta_parts.append(exp.location)

                    if meta_parts:
                        role_date_para.add_run(f"   {', '.join(meta_parts)}")
                        meta_run = role_date_para.runs[-1]
                        meta_run.font.size = Pt(9)
                        meta_run.font.color.rgb = RGBColor(128, 128, 128)

                # Minimal bullet points
                if exp.description:
                    description_bullets = self._parse_description_into_bullets(exp.description)
                    for bullet in description_bullets:
                        bullet_para = doc.add_paragraph()
                        bullet_run = bullet_para.add_run(bullet)
                        bullet_run.font.size = Pt(10)
                        bullet_run.font.name = 'Helvetica'
                        bullet_run.font.color.rgb = RGBColor(89, 89, 89)
                        bullet_para.paragraph_format.left_indent = Inches(0.2)

                        # Add subtle bullet
                        bullet_marker = bullet_para.add_run("•")
                        bullet_marker.font.color.rgb = RGBColor(0, 184, 148)
                        bullet_marker.font.size = Pt(12)

        # ===== SKILLS =====
        if profile.skills:
            doc.add_paragraph()

            skills_header = doc.add_paragraph()
            skills_header_run = skills_header.add_run("SKILLS")
            skills_header_run.font.size = Pt(16)
            skills_header_run.bold = True
            skills_header_run.font.name = 'Helvetica'
            skills_header_run.font.color.rgb = RGBColor(0, 184, 148)

            # Clean skills layout
            skills_text = ", ".join([skill.name for skill in profile.skills])
            skills_para = doc.add_paragraph()
            skills_run = skills_para.add_run(skills_text)
            skills_run.font.size = Pt(10)
            skills_run.font.name = 'Helvetica'
            skills_run.font.color.rgb = RGBColor(64, 64, 64)

    async def _apply_tech_template(self, doc: Document, profile: ProfileSchema, tailoring: TailorResponse):
        """Apply tech-focused template - monospace fonts, code-friendly formatting"""
        # Set tech margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.7)
            section.bottom_margin = Inches(0.7)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)

        # ===== TECH HEADER =====
        # Name in monospace font
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(f"{profile.basics.firstName} {profile.basics.lastName}")
        name_run.font.size = Pt(18)
        name_run.bold = True
        name_run.font.name = 'Consolas'
        name_run.font.color.rgb = RGBColor(0, 43, 54)  # Dark teal

        # Tech contact info
        contact_parts = []
        if profile.basics.email:
            contact_parts.append(f"email: {profile.basics.email}")
        if profile.basics.phone:
            contact_parts.append(f"phone: {profile.basics.phone}")
        if profile.basics.location:
            contact_parts.append(f"location: {profile.basics.location}")

        if profile.basics.links:
            for link in profile.basics.links:
                if link.url:
                    if link.type == "GitHub":
                        contact_parts.append(f"github: {link.url}")
                    elif link.type == "LinkedIn":
                        contact_parts.append(f"linkedin: {link.url}")
                    elif link.type == "Portfolio":
                        contact_parts.append(f"portfolio: {link.url}")

        if contact_parts:
            contact_para = doc.add_paragraph("\n".join(contact_parts))
            contact_run = contact_para.runs[0]
            contact_run.font.size = Pt(9)
            contact_run.font.name = 'Consolas'
            contact_run.font.color.rgb = RGBColor(88, 110, 117)

        # Code-style separator
        separator = doc.add_paragraph()
        separator_run = separator.add_run("#" * 50)
        separator_run.font.size = Pt(8)
        separator_run.font.name = 'Consolas'
        separator_run.font.color.rgb = RGBColor(38, 139, 210)
        separator.paragraph_format.space_before = Pt(12)
        separator.paragraph_format.space_after = Pt(12)

        # ===== TECHNICAL EXPERIENCE =====
        if profile.work_experience:
            exp_header = doc.add_paragraph()
            exp_header_run = exp_header.add_run("## TECHNICAL EXPERIENCE")
            exp_header_run.font.size = Pt(14)
            exp_header_run.bold = True
            exp_header_run.font.name = 'Consolas'
            exp_header_run.font.color.rgb = RGBColor(38, 139, 210)

            for exp in profile.work_experience:
                doc.add_paragraph()

                # Company and role
                company_role = f"{exp.company} | {exp.title}"
                company_para = doc.add_paragraph()
                company_run = company_para.add_run(company_role)
                company_run.font.size = Pt(12)
                company_run.bold = True
                company_run.font.name = 'Consolas'
                company_run.font.color.rgb = RGBColor(0, 43, 54)

                # Timeline and location
                if exp.startDate or exp.location:
                    timeline = []
                    if exp.startDate:
                        date_str = exp.startDate
                        if exp.endDate:
                            date_str += f" - {exp.endDate}"
                        else:
                            date_str += " - Present"
                        timeline.append(date_str)
                    if exp.location:
                        timeline.append(exp.location)

                    if timeline:
                        timeline_para = doc.add_paragraph()
                        timeline_run = timeline_para.add_run(f"[{', '.join(timeline)}]")
                        timeline_run.font.size = Pt(9)
                        timeline_run.font.name = 'Consolas'
                        timeline_run.font.color.rgb = RGBColor(88, 110, 117)

                # Technical bullet points
                if exp.description:
                    description_bullets = self._parse_description_into_bullets(exp.description)
                    for bullet in description_bullets:
                        bullet_para = doc.add_paragraph()
                        bullet_run = bullet_para.add_run(f"• {bullet}")
                        bullet_run.font.size = Pt(10)
                        bullet_run.font.name = 'Consolas'
                        bullet_run.font.color.rgb = RGBColor(101, 123, 131)
                        bullet_para.paragraph_format.left_indent = Inches(0.25)

                # Tech stack
                if exp.technologies and len(exp.technologies) > 0:
                    tech_para = doc.add_paragraph()
                    tech_run = tech_para.add_run(f"Tech Stack: {', '.join(exp.technologies)}")
                    tech_run.font.size = Pt(9)
                    tech_run.font.name = 'Consolas'
                    tech_run.font.color.rgb = RGBColor(38, 139, 210)

        # ===== SKILLS MATRIX =====
        if profile.skills:
            doc.add_paragraph()

            skills_header = doc.add_paragraph()
            skills_header_run = skills_header.add_run("## SKILLS MATRIX")
            skills_header_run.font.size = Pt(14)
            skills_header_run.bold = True
            skills_header_run.font.name = 'Consolas'
            skills_header_run.font.color.rgb = RGBColor(38, 139, 210)

            # Group skills by category
            skills_by_category = {}
            for skill in profile.skills:
                category = skill.category or 'General'
                if category not in skills_by_category:
                    skills_by_category[category] = []
                level_indicator = ""
                if skill.level:
                    if skill.level == "expert":
                        level_indicator = "●●●"
                    elif skill.level == "advanced":
                        level_indicator = "●●○"
                    elif skill.level == "intermediate":
                        level_indicator = "●○○"
                    else:
                        level_indicator = "○○○"
                skills_by_category[category].append(f"{skill.name} {level_indicator}")

            for category, skills in skills_by_category.items():
                cat_para = doc.add_paragraph()
                cat_run = cat_para.add_run(f"{category}:")
                cat_run.font.size = Pt(11)
                cat_run.bold = True
                cat_run.font.name = 'Consolas'
                cat_run.font.color.rgb = RGBColor(0, 43, 54)

                skills_para = doc.add_paragraph()
                skills_run = skills_para.add_run(", ".join(skills))
                skills_run.font.size = Pt(10)
                skills_run.font.name = 'Consolas'
                skills_run.font.color.rgb = RGBColor(101, 123, 131)

    async def _apply_academic_template(self, doc: Document, profile: ProfileSchema, tailoring: TailorResponse):
        """Apply academic formal template - emphasis on credentials and research"""
        # Set academic margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)

        # ===== ACADEMIC HEADER =====
        # Formal name and credentials
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(f"{profile.basics.firstName} {profile.basics.lastName}")
        name_run.font.size = Pt(18)
        name_run.bold = True
        name_run.font.name = 'Garamond'
        name_run.font.color.rgb = RGBColor(0, 0, 0)
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Academic title/position
        if profile.basics.headline:
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(profile.basics.headline)
            title_run.font.size = Pt(12)
            title_run.font.name = 'Garamond'
            title_run.font.color.rgb = RGBColor(89, 89, 89)
            title_run.italic = True
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Professional contact
        contact_parts = []
        if profile.basics.email:
            contact_parts.append(f"Email: {profile.basics.email}")
        if profile.basics.phone:
            contact_parts.append(f"Phone: {profile.basics.phone}")
        if profile.basics.location:
            contact_parts.append(f"Location: {profile.basics.location}")

        if contact_parts:
            contact_para = doc.add_paragraph(" • ".join(contact_parts))
            contact_run = contact_para.runs[0]
            contact_run.font.size = Pt(10)
            contact_run.font.name = 'Garamond'
            contact_run.font.color.rgb = RGBColor(89, 89, 89)
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Academic separator
        separator = doc.add_paragraph()
        separator_run = separator.add_run("❖ ❖ ❖")
        separator_run.font.size = Pt(12)
        separator_run.font.name = 'Garamond'
        separator_run.font.color.rgb = RGBColor(139, 69, 19)  # Saddle brown
        separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
        separator.paragraph_format.space_before = Pt(12)
        separator.paragraph_format.space_after = Pt(12)

        # ===== PROFESSIONAL EXPERIENCE =====
        if profile.work_experience:
            exp_header = doc.add_paragraph()
            exp_header_run = exp_header.add_run("PROFESSIONAL EXPERIENCE")
            exp_header_run.font.size = Pt(14)
            exp_header_run.bold = True
            exp_header_run.font.name = 'Garamond'
            exp_header_run.font.color.rgb = RGBColor(139, 69, 19)

            for exp in profile.work_experience:
                doc.add_paragraph()

                # Institution/Company
                company_para = doc.add_paragraph()
                company_run = company_para.add_run(exp.company)
                company_run.font.size = Pt(12)
                company_run.bold = True
                company_run.font.name = 'Garamond'
                company_run.font.color.rgb = RGBColor(0, 0, 0)

                # Position and dates
                position_para = doc.add_paragraph()
                position_run = position_para.add_run(exp.title)
                position_run.font.size = Pt(11)
                position_run.font.name = 'Garamond'
                position_run.font.color.rgb = RGBColor(64, 64, 64)

                if exp.startDate or exp.location:
                    meta_parts = []
                    if exp.startDate:
                        date_str = exp.startDate
                        if exp.endDate:
                            date_str += f" – {exp.endDate}"
                        else:
                            date_str += " – Present"
                        meta_parts.append(date_str)
                    if exp.location:
                        meta_parts.append(exp.location)

                    if meta_parts:
                        position_para.add_run(f", {', '.join(meta_parts)}")
                        meta_run = position_para.runs[-1]
                        meta_run.font.size = Pt(10)
                        meta_run.font.color.rgb = RGBColor(89, 89, 89)

                # Academic-style bullet points
                if exp.description:
                    description_bullets = self._parse_description_into_bullets(exp.description)
                    for bullet in description_bullets:
                        bullet_para = doc.add_paragraph()
                        bullet_run = bullet_para.add_run(f"— {bullet}")
                        bullet_run.font.size = Pt(10)
                        bullet_run.font.name = 'Garamond'
                        bullet_run.font.color.rgb = RGBColor(64, 64, 64)
                        bullet_para.paragraph_format.left_indent = Inches(0.25)

        # ===== ACADEMIC BACKGROUND =====
        if profile.education:
            doc.add_paragraph()

            edu_header = doc.add_paragraph()
            edu_header_run = edu_header.add_run("ACADEMIC BACKGROUND")
            edu_header_run.font.size = Pt(14)
            edu_header_run.bold = True
            edu_header_run.font.name = 'Garamond'
            edu_header_run.font.color.rgb = RGBColor(139, 69, 19)

            for edu in profile.education:
                doc.add_paragraph()

                # Institution
                school_para = doc.add_paragraph()
                school_run = school_para.add_run(edu.school)
                school_run.font.size = Pt(12)
                school_run.bold = True
                school_run.font.name = 'Garamond'
                school_run.font.color.rgb = RGBColor(0, 0, 0)

                # Degree and field
                degree_para = doc.add_paragraph()
                degree_text = edu.degree
                if edu.field:
                    degree_text += f" in {edu.field}"
                degree_run = degree_para.add_run(degree_text)
                degree_run.font.size = Pt(11)
                degree_run.font.name = 'Garamond'
                degree_run.font.color.rgb = RGBColor(64, 64, 64)

                # Academic honors and dates
                if edu.endDate or (edu.honors and len(edu.honors) > 0):
                    honors_parts = []
                    if edu.endDate:
                        honors_parts.append(f"{edu.endDate}")
                    if edu.honors and len(edu.honors) > 0:
                        honors_parts.append(f"Honors: {', '.join(edu.honors)}")

                    if honors_parts:
                        honors_para = doc.add_paragraph()
                        honors_run = honors_para.add_run(", ".join(honors_parts))
                        honors_run.font.size = Pt(10)
                        honors_run.font.name = 'Garamond'
                        honors_run.font.color.rgb = RGBColor(89, 89, 89)

        # ===== PUBLICATIONS & RESEARCH =====
        if profile.certifications or profile.awards:
            doc.add_paragraph()

            research_header = doc.add_paragraph()
            research_header_run = research_header.add_run("CERTIFICATIONS & AWARDS")
            research_header_run.font.size = Pt(14)
            research_header_run.bold = True
            research_header_run.font.name = 'Garamond'
            research_header_run.font.color.rgb = RGBColor(139, 69, 19)

            # Certifications
            if profile.certifications:
                for cert in profile.certifications:
                    cert_para = doc.add_paragraph()
                    cert_run = cert_para.add_run(f"• {cert.name}")
                    cert_run.font.size = Pt(10)
                    cert_run.font.name = 'Garamond'
                    cert_run.font.color.rgb = RGBColor(64, 64, 64)

                    if cert.issuer or cert.date:
                        issuer_parts = []
                        if cert.issuer:
                            issuer_parts.append(cert.issuer)
                        if cert.date:
                            issuer_parts.append(cert.date)
                        if issuer_parts:
                            cert_para.add_run(f", {', '.join(issuer_parts)}")
                            issuer_run = cert_para.runs[-1]
                            issuer_run.font.size = Pt(9)
                            issuer_run.font.color.rgb = RGBColor(89, 89, 89)

            # Awards
            if profile.awards:
                for award in profile.awards:
                    award_para = doc.add_paragraph()
                    award_run = award_para.add_run(f"• {award.name}")
                    award_run.font.size = Pt(10)
                    award_run.font.name = 'Garamond'
                    award_run.font.color.rgb = RGBColor(64, 64, 64)

                    if award.issuer or award.date:
                        award_parts = []
                        if award.issuer:
                            award_parts.append(award.issuer)
                        if award.date:
                            award_parts.append(award.date)
                        if award_parts:
                            award_para.add_run(f", {', '.join(award_parts)}")
                            award_meta_run = award_para.runs[-1]
                            award_meta_run.font.size = Pt(9)
                            award_meta_run.font.color.rgb = RGBColor(89, 89, 89)

    async def _generate_cover_letter_docx(
        self,
        profile: ProfileSchema,
        tailoring: TailorResponse,
        original_profile_id: Optional[str]
    ) -> Optional[str]:
        """Generate cover letter DOCX"""
        if not tailoring.cover_letter_text:
            app_logger.log_info("No cover letter text available, skipping cover letter generation")
            return None

        file_id = str(uuid.uuid4())
        file_path = self.upload_dir / f"{file_id}_cover_letter.docx"

        try:
            doc = Document()

            # Set margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)

            # Header - Contact info
            header = doc.add_paragraph()
            header.alignment = 1  # Center
            run = header.add_run(f"{profile.basics.firstName} {profile.basics.lastName}")
            run.bold = True
            run.font.size = Pt(14)

            # Contact info
            contact_info = []
            if profile.basics.email:
                contact_info.append(profile.basics.email)
            if profile.basics.phone:
                contact_info.append(profile.basics.phone)
            if profile.basics.location:
                contact_info.append(profile.basics.location)

            if contact_info:
                contact_para = doc.add_paragraph()
                contact_para.alignment = 1  # Center
                contact_run = contact_para.add_run(" | ".join(contact_info))
                contact_run.font.size = Pt(10)

            # Add spacing
            doc.add_paragraph()

            # Date
            import datetime
            date_para = doc.add_paragraph()
            date_run = date_para.add_run(datetime.date.today().strftime("%B %d, %Y"))
            date_run.font.size = Pt(10)

            # Add spacing
            doc.add_paragraph()

            # Hiring manager greeting (generic)
            greeting = doc.add_paragraph()
            greeting_run = greeting.add_run("Dear Hiring Manager,")
            greeting_run.font.size = Pt(11)

            # Cover letter content - clean and format properly
            content = tailoring.cover_letter_text.strip()

            # Split into paragraphs and clean up
            paragraphs = [p for p in content.split('\n\n') if p.strip()]

            for paragraph in paragraphs:
                # Skip if this is a duplicate greeting or closing
                clean_para = paragraph.strip().lower()
                skip_keywords = ['dear hiring manager', 'sincerely', 'best regards']
                if any(keyword in clean_para for keyword in skip_keywords) or \
                   clean_para == f"{profile.basics.firstName} {profile.basics.lastName}".lower():
                    continue

                para = doc.add_paragraph()
                run = para.add_run(clean_para)
                run.font.size = Pt(11)

                # Add spacing between paragraphs
                doc.add_paragraph()

            # Closing
            doc.add_paragraph()
            closing = doc.add_paragraph()
            closing_run = closing.add_run("Sincerely,")
            closing_run.font.size = Pt(11)

            # Name
            doc.add_paragraph()
            name_para = doc.add_paragraph()
            name_run = name_para.add_run(f"{profile.basics.firstName} {profile.basics.lastName}")
            name_run.bold = True
            name_run.font.size = Pt(11)

            # Save document
            doc.save(str(file_path))
            app_logger.log_info(f"Cover letter DOCX generated: {file_id}")
            return file_id

        except Exception as e:
            app_logger.log_error(f"Cover letter DOCX generation error: {e}")
            return None

    async def _generate_application_package_zip(
        self,
        profile: ProfileSchema,
        tailoring: TailorResponse,
        job_posting: Any,
        docx_resume_path: Optional[str],
        pdf_resume_path: Optional[str],
        cover_letter_path: Optional[str],
        original_profile_id: Optional[str]
    ) -> Optional[str]:
        """Generate complete application package ZIP (containing resume DOCX, PDF, and cover letter)"""
        file_id = str(uuid.uuid4())
        zip_path = self.upload_dir / f"{file_id}_application_package.zip"

        try:
            with zipfile.ZipFile(str(zip_path), 'w', zipfile.ZIP_DEFLATED) as zip_file:
                # Add resume DOCX if available
                if docx_resume_path and Path(docx_resume_path).exists():
                    zip_file.write(docx_resume_path, f"Tailored_Resume_{profile.basics.firstName}_{profile.basics.lastName}.docx")

                # Add resume PDF if available
                if pdf_resume_path and Path(pdf_resume_path).exists():
                    zip_file.write(pdf_resume_path, f"Tailored_Resume_{profile.basics.firstName}_{profile.basics.lastName}.pdf")

                # Add cover letter if available
                if cover_letter_path and Path(cover_letter_path).exists():
                    zip_file.write(cover_letter_path, f"Cover_Letter_{profile.basics.firstName}_{profile.basics.lastName}.docx")

            app_logger.log_info(f"Application package ZIP generated: {file_id}")
            return file_id

        except Exception as e:
            app_logger.log_error(f"Application package ZIP generation error: {e}")
            return None

document_generator = DocumentGenerator()
