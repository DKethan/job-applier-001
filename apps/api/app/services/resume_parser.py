import json
import re
import traceback
from docx import Document
from typing import Dict, Any
from openai import OpenAI
from app.config import settings
from app.schemas.profile import ProfileSchema, Basics
from pydantic import ValidationError
from app.utils.app_logger import app_logger

# Try to import Docling and PyPDF2 for fallbacks
try:
    from docling.document_converter import DocumentConverter
    DOCLING_AVAILABLE = True
except ImportError:
    DOCLING_AVAILABLE = False

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

# Logger is now app_logger from imports above


def convert_pdf_to_markdown(pdf_path: str, debug: bool = False) -> str:
    """
    Convert PDF to Markdown using Docling.

    Args:
        pdf_path: Path to PDF file
        debug: Enable debug logging

    Returns:
        Markdown text representation of PDF
    """
    try:
        if debug:
            app_logger.log_info(f"[convert_pdf_to_markdown] Converting PDF to Markdown: {pdf_path}")

        # Suppress docling INFO logs
        import logging as logging_module
        docling_logger = logging_module.getLogger("docling")
        original_level = docling_logger.level
        docling_logger.setLevel(logging_module.WARNING)

        # Create Docling converter
        converter = DocumentConverter()

        # Convert PDF to document
        result = converter.convert(pdf_path)

        # Export to Markdown
        markdown_text = result.document.export_to_markdown()

        # Restore original logging level
        docling_logger.setLevel(original_level)

        if debug:
            app_logger.log_info(f"[convert_pdf_to_markdown] Successfully converted PDF to Markdown ({len(markdown_text)} chars)")

        return markdown_text

    except ImportError:
        error_msg = "Docling not installed. Please install: pip install docling"
        app_logger.log_error(f"[convert_pdf_to_markdown] {error_msg}")
        return ""
    except Exception as e:
        error_msg = f"Docling conversion failed: {str(e)}"
        app_logger.log_error(f"[convert_pdf_to_markdown] {error_msg}")

        # Try a simpler fallback approach
        try:
            if debug:
                app_logger.log_info("[convert_pdf_to_markdown] Trying basic PDF text extraction fallback")

            # Try to use PyPDF2 if available
            if PYPDF2_AVAILABLE and PyPDF2:
                try:
                    with open(pdf_path, 'rb') as file:
                        pdf_reader = PyPDF2.PdfReader(file)
                        text = ""
                        for page in pdf_reader.pages:
                            text += page.extract_text() + "\n"
                        if len(text.strip()) >= 50:
                            app_logger.log_info(f"[convert_pdf_to_markdown] Fallback extraction successful ({len(text)} chars)")
                            return text
                except Exception as e:
                    app_logger.log_warning(f"[convert_pdf_to_markdown] PyPDF2 extraction failed: {str(e)}")
                    pass

            # Try basic file reading as last resort (for text-based PDFs)
            try:
                with open(pdf_path, 'rb') as f:
                    content = f.read()
                    # Look for text content in PDF (very basic)
                    text_content = content.decode('utf-8', errors='ignore')
                    # Extract text between parentheses and other common patterns
                    text_matches = re.findall(r'\(([^)]+)\)', text_content)
                    if text_matches:
                        extracted = ' '.join(text_matches)
                        if len(extracted.strip()) >= 20:
                            app_logger.log_info(f"[convert_pdf_to_markdown] Basic text extraction successful ({len(extracted)} chars)")
                            return extracted
            except Exception:
                pass

        except Exception as fallback_e:
            app_logger.log_error(f"[convert_pdf_to_markdown] Fallback also failed: {str(fallback_e)}")

        if debug:
            app_logger.log_error(f"[convert_pdf_to_markdown] Traceback: {traceback.format_exc()}")
        return ""


class ResumeParser:
    def __init__(self):
        self.client = OpenAI(api_key=settings.openai_api_key) if settings.openai_api_key else None

    def extract_text(self, file_path: str, content_type: str) -> str:
        """Extract text from PDF or DOCX"""
        if content_type == "application/pdf":
            return self._extract_pdf_text(file_path)
        elif content_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
            return self._extract_docx_text(file_path)
        else:
            raise ValueError(f"Unsupported file type: {content_type}")

    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF using Docling (returns markdown)"""
        try:
            markdown_text = convert_pdf_to_markdown(file_path, debug=True)
            if markdown_text:
                return markdown_text
            else:
                raise ValueError("Docling conversion returned empty result")
        except Exception as e:
            app_logger.log_error(f"Error extracting PDF text with Docling: {e}")
            raise

    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX using python-docx"""
        text = ""
        try:
            doc = Document(file_path)
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + "\n"
        except Exception as e:
            app_logger.log_error(f"Error extracting DOCX text: {e}")
            raise
        return text

    def _post_process_llm_response(self, profile_data: dict) -> dict:
        """Post-process LLM response to fix common structural issues"""
        # Ensure required top-level keys exist
        required_keys = ["basics", "work_experience", "education", "projects", "skills"]
        for key in required_keys:
            if key not in profile_data:
                if key == "basics":
                    profile_data[key] = {
                        "firstName": "",
                        "lastName": "",
                        "email": None,
                        "links": []
                    }
                else:
                    profile_data[key] = []

        # Fix basics.links - convert strings to proper link objects
        if "basics" in profile_data and isinstance(profile_data["basics"], dict):
            if "links" in profile_data["basics"] and isinstance(profile_data["basics"]["links"], list):
                fixed_links = []
                for link in profile_data["basics"]["links"]:
                    if isinstance(link, str):
                        # Convert string to link object
                        fixed_links.append({
                            "type": link,
                            "url": "",
                            "label": link
                        })
                    elif isinstance(link, dict):
                        fixed_links.append(link)
                profile_data["basics"]["links"] = fixed_links

        # Fix work_experience - convert jobTitle to title
        if "work_experience" in profile_data and isinstance(profile_data["work_experience"], list):
            for exp in profile_data["work_experience"]:
                if isinstance(exp, dict):
                    if "jobTitle" in exp and "title" not in exp:
                        exp["title"] = exp["jobTitle"]
                        del exp["jobTitle"]

        # Fix projects - convert tech strings to arrays
        if "projects" in profile_data and isinstance(profile_data["projects"], list):
            for project in profile_data["projects"]:
                if isinstance(project, dict):
                    if "tech" in project and isinstance(project["tech"], str):
                        # Split comma-separated string into array
                        project["tech"] = [tech.strip() for tech in project["tech"].split(",")]

        # Fix skills - convert strings to skill objects
        if "skills" in profile_data and isinstance(profile_data["skills"], list):
            fixed_skills = []
            for skill in profile_data["skills"]:
                if isinstance(skill, str):
                    # Convert string to skill object
                    fixed_skills.append({
                        "name": skill,
                        "level": None
                    })
                elif isinstance(skill, dict):
                    fixed_skills.append(skill)
            profile_data["skills"] = fixed_skills

        # Fix education structure - convert 'institution' to 'school' if needed
        if "education" in profile_data and isinstance(profile_data["education"], list):
            for edu in profile_data["education"]:
                if isinstance(edu, dict):
                    if "institution" in edu and "school" not in edu:
                        edu["school"] = edu["institution"]
                    # Handle GPA/CGPA as strings or numbers
                    if "cgpa" in edu and "gpa" not in edu:
                        edu["gpa"] = edu["cgpa"]
                    # Convert GPA strings like "8.78 / 10.0" to numbers
                    if "gpa" in edu and isinstance(edu["gpa"], str):
                        try:
                            # Extract first number from strings like "8.78 / 10.0"
                            gpa_match = re.match(r'(\d+\.?\d*)', edu["gpa"])
                            if gpa_match:
                                edu["gpa"] = float(gpa_match.group(1))
                        except (ValueError, AttributeError):
                            pass  # Keep as string if can't parse

        return profile_data

    def parse_with_llm(self, resume_text: str) -> ProfileSchema:
        """Parse resume text using LLM with structured output"""
        if not self.client:
            raise ValueError("OpenAI client not initialized. Check OPENAI_API_KEY.")

        # Create JSON schema for Profile
        profile_schema = {
            "type": "object",
            "properties": {
                "basics": {
                    "type": "object",
                    "properties": {
                        "firstName": {"type": "string"},
                        "lastName": {"type": "string"},
                        "email": {"type": ["string", "null"], "format": "email"},
                        "phone": {"type": "string"},
                        "location": {"type": "string"},
                        "headline": {"type": "string"},
                        "summary": {"type": "string"},
                        "languages": {
                            "type": "array",
                            "items": {"type": "string"}
                        },
                        "links": {
                            "type": "array",
                            "items": {
                                "type": "object",
                                "properties": {
                                    "type": {"type": "string", "description": "Type of link (Portfolio, LinkedIn, GitHub, etc.)"},
                                    "url": {"type": "string", "description": "URL of the link"},
                                    "label": {"type": "string", "description": "Display label for the link"}
                                },
                                "required": ["type"]
                            }
                        }
                    },
                    "required": ["firstName", "lastName"]
                },
                "work_experience": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "company": {"type": "string"},
                            "title": {"type": "string"},
                            "jobTitle": {"type": "string"},
                            "startDate": {"type": "string"},
                            "endDate": {"type": "string"},
                            "location": {"type": "string"},
                            "employmentType": {"type": "string"},
                            "description": {"type": "string"},
                            "technologies": {"type": "array", "items": {"type": "string"}},
                            "achievements": {"type": "array", "items": {"type": "string"}},
                            "bullets": {"type": "array", "items": {"type": "string"}}
                        },
                        "required": ["company", "title", "startDate", "bullets"]
                    }
                },
                "education": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "school": {"type": "string"},
                            "institution": {"type": "string"},
                            "degree": {"type": "string"},
                            "field": {"type": "string"},
                            "startDate": {"type": "string"},
                            "endDate": {"type": "string"},
                            "gpa": {"type": ["number", "string"]},
                            "cgpa": {"type": ["number", "string"]},
                            "grade": {"type": "string"},
                            "honors": {"type": "array", "items": {"type": "string"}},
                            "relevantCoursework": {"type": "array", "items": {"type": "string"}},
                            "activities": {"type": "array", "items": {"type": "string"}},
                            "thesis": {"type": "string"},
                            "bullets": {"type": "array", "items": {"type": "string"}}
                        },
                        "required": ["school", "institution", "degree"]
                    }
                },
                "projects": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "name": {"type": "string"},
                            "link": {"type": "string"},
                            "description": {"type": "string"},
                            "startDate": {"type": "string"},
                            "endDate": {"type": "string"},
                            "role": {"type": "string"},
                            "teamSize": {"type": "number"},
                            "status": {"type": "string"},
                            "outcomes": {"type": "array", "items": {"type": "string"}},
                            "tech": {"type": "array", "items": {"type": "string"}, "description": "Array of technologies used"},
                            "bullets": {"type": "array", "items": {"type": "string"}}
                        },
                        "required": ["name", "description"]
                    }
                },
                "skills": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "name": {"type": "string", "description": "Name of the skill"},
                            "category": {"type": "string", "description": "Category (Programming, Tools, Soft Skills, etc.)"},
                            "level": {"type": "string", "enum": ["beginner", "intermediate", "advanced", "expert"], "description": "Proficiency level"},
                            "yearsExperience": {"type": "number", "description": "Years of experience"}
                        },
                        "required": ["name"]
                    },
                    "description": "Array of skill objects, each with name and optional level"
                },
                "languages": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "name": {"type": "string"},
                            "proficiency": {"type": "string", "enum": ["native", "fluent", "conversational", "basic"]}
                        },
                        "required": ["name"]
                    }
                },
                "certifications": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "name": {"type": "string"},
                            "issuer": {"type": "string"},
                            "date": {"type": "string"},
                            "expiryDate": {"type": "string"},
                            "credentialId": {"type": "string"},
                            "url": {"type": "string"},
                            "description": {"type": "string"},
                            "skills": {"type": "array", "items": {"type": "string"}},
                            "verification": {"type": "string"}
                        },
                        "required": ["name"]
                    }
                },
                "awards": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "name": {"type": "string"},
                            "issuer": {"type": "string"},
                            "date": {"type": "string"},
                            "category": {"type": "string"},
                            "description": {"type": "string"},
                            "recognition": {"type": "string"}
                        },
                        "required": ["name"]
                    }
                },
                "preferences": {
                    "type": "object",
                    "properties": {
                        "visaStatus": {"type": "string"},
                        "workAuth": {"type": "string"},
                        "relocation": {"type": "boolean"},
                        "remote": {"type": "boolean"},
                        "salary": {"type": "string"},
                        "availability": {"type": "string"}
                    }
                }
            },
            "required": ["basics", "work_experience", "education", "projects", "skills"]
        }

        prompt = f"""Extract structured profile information from the following resume text.
The resume is provided in markdown format, which preserves document structure like headings, lists, and formatting.

IMPORTANT EXTRACTION RULES:
- For WORK EXPERIENCE: Extract detailed job descriptions and key achievements from the resume text. Look for paragraphs describing responsibilities, accomplishments, and impact.
- For PROJECTS: Extract detailed descriptions of what each project was about, the problems solved, and outcomes achieved.
- For EDUCATION: Extract detailed information about coursework, activities, honors, and academic achievements.
- Extract ALL available information - don't leave fields empty if the data exists in the resume.

CRITICAL: Return a JSON object with EXACTLY these top-level keys and formats:

- basics: {{
    "firstName": "string",
    "lastName": "string",
    "email": "string or null",
    "phone": "string or null",
    "location": "string or null",
    "headline": "string or null",
    "summary": "string or null",
    "languages": ["string"] or null,
    "links": [ {{"type": "string", "url": "string", "label": "string"}} ]
  }}

- work_experience: [
    {{
      "company": "string",
      "title": "string",
      "startDate": "string",
      "endDate": "string or null",
      "location": "string or null",
      "employmentType": "full-time|part-time|contract|freelance|internship or null",
      "description": "DETAILED job description from resume - extract full paragraphs describing the role, responsibilities, and scope of work",
      "technologies": ["string - extract all technical skills, tools, and technologies mentioned for this role"],
      "achievements": ["string - extract specific achievements, metrics, and quantifiable results from the resume"],
      "bullets": ["string - extract detailed bullet points describing responsibilities and accomplishments"]
    }}
  ]

- education: [
    {{
      "school": "string",
      "degree": "string",
      "field": "string or null",
      "startDate": "string or null",
      "endDate": "string or null",
      "gpa": "number or null",
      "grade": "string or null",
      "honors": ["string"] or null,
      "relevantCoursework": ["string"] or null,
      "activities": ["string"] or null,
      "thesis": "string or null",
      "bullets": ["string"] or null
    }}
  ]

- projects: [
    {{
      "name": "string",
      "link": "string or null",
      "description": "DETAILED project description - extract what the project was about, the problem it solved, and its purpose",
      "startDate": "string or null",
      "endDate": "string or null",
      "role": "string or null - your specific role/contribution to the project",
      "teamSize": "number or null",
      "status": "completed|in-progress|on-hold or null",
      "outcomes": ["string - extract measurable results, impact, and achievements of the project"],
      "tech": ["string - all technologies, frameworks, and tools used in the project"],
      "bullets": ["string - detailed contributions and responsibilities in the project"]
    }}
  ]

- skills: [
    {{
      "name": "string",
      "category": "Programming|Tools|Soft Skills etc. or null",
      "level": "beginner|intermediate|advanced|expert or null",
      "yearsExperience": "number or null"
    }}
  ]

- languages: [
    {{
      "name": "string",
      "proficiency": "native|fluent|conversational|basic or null"
    }}
  ] or null

- certifications: [
    {{
      "name": "string",
      "issuer": "string or null",
      "date": "string or null",
      "expiryDate": "string or null",
      "credentialId": "string or null",
      "url": "string or null",
      "description": "string or null",
      "skills": ["string"] or null,
      "verification": "string or null"
    }}
  ] or null

- awards: [
    {{
      "name": "string",
      "issuer": "string or null",
      "date": "string or null",
      "category": "Academic|Professional|Leadership etc. or null",
      "description": "string or null",
      "recognition": "string or null"
    }}
  ] or null

- preferences: {{
    "visaStatus": "string or null",
    "workAuth": "string or null",
    "relocation": "boolean or null",
    "remote": "boolean or null",
    "salary": "string or null",
    "availability": "string or null"
  }} or null

IMPORTANT:
- READ THE ENTIRE RESUME CAREFULLY and extract ALL detailed information available.
- For job descriptions: Look for paragraphs that describe what the person did, their responsibilities, and the scope of their work.
- For achievements: Extract specific metrics, results, and accomplishments mentioned in the resume.
- For projects: Extract detailed descriptions of what each project accomplished and the person's contributions.
- Do NOT return strings for arrays like technologies, achievements, honors, etc. Always use arrays.
- Do NOT fabricate information - only extract what exists in the resume.
- If a section has no data, use null or empty array.
- For dates, use formats like "2023-01" or "January 2023" or "2023".
- Extract ALL available information from the resume - be thorough!

Resume text (markdown formatted):
{resume_text}

Return only the JSON object, no other text or explanations."""

        try:
            app_logger.log_info(f"Sending resume text to LLM ({len(resume_text)} chars)")
            response = self.client.chat.completions.create(
                model=settings.openai_model,
                messages=[
                    {"role": "system", "content": "You are a resume parser that extracts structured data into JSON format. Always return valid JSON only."},
                    {"role": "user", "content": prompt}
                ],
                response_format={"type": "json_object"},
                temperature=0.1,
            )

            content = response.choices[0].message.content
            app_logger.log_info(f"LLM response received ({len(content) if content else 0} chars)")
            if not content:
                raise ValueError("LLM returned empty response")

            profile_data = json.loads(content)

            # Post-process the response to fix common issues
            app_logger.log_info(f"Raw LLM response: {profile_data}")
            profile_data = self._post_process_llm_response(profile_data)
            app_logger.log_info(f"Post-processed response: {profile_data}")

            # Validate against Pydantic schema
            try:
                profile = ProfileSchema(**profile_data)
                return profile
            except Exception as e:
                app_logger.log_error(f"Profile validation failed: {e}")
                # Return minimal valid profile
                return ProfileSchema(
                    basics=Basics(
                        firstName="",
                        lastName="",
                        email=None,
                        links=[]
                    ),
                    work_experience=[],
                    education=[],
                    projects=[],
                    skills=[]
                )

        except json.JSONDecodeError as e:
            app_logger.log_error(f"Invalid JSON from LLM: {e}")
            app_logger.log_error(f"Raw LLM response: {content}")
            raise ValueError(f"Failed to parse LLM response as JSON: {e}")
        except ValidationError as e:
            app_logger.log_error(f"Profile validation error: {e}")
            app_logger.log_error(f"LLM response that failed validation: {profile_data}")
            raise ValueError(f"Invalid profile structure: {e}")
        except Exception as e:
            app_logger.log_error(f"Error calling LLM: {e}")
            app_logger.log_error(f"LLM content: {content}")
            raise


resume_parser = ResumeParser()
